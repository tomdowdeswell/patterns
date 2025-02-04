import requests
import json
from youtube_transcript_api import YouTubeTranscriptApi
import re
from docx import Document
import requests
from bs4 import BeautifulSoup

# Define the allowlist of characters
ALLOWLIST_PATTERN = re.compile(r"^[a-zA-Z0-9\s.,;:!?\-]+$")


def get_title(video_id):
    # YouTube Video URL
    url = f'https://www.youtube.com/watch?v={video_id}'

    # Extracting HTML Code of the Video Page:
    response = requests.get(url)
    html_content = response.text

    # Processing the HTML Code with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Extracting <title> tag's content
    title_tag = soup.find('meta', property='og:title')
    video_title = title_tag['content'] if title_tag else 'Title not found'
    
    return video_title
    
def sanitize_filename(filename):
    # Replace invalid characters with an underscore
    return re.sub(r'[\\/*?:"<>|]', '_', filename)

# Sanitize the content, sort of. Prompt injection is the main threat so this isn't a huge deal
def sanitize_content(content):
    return "".join(char for char in content if ALLOWLIST_PATTERN.match(char))


# Pull the URL content's from the GitHub repo
def fetch_content_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        sanitized_content = sanitize_content(response.text)
        return sanitized_content
    except requests.RequestException as e:
        return str(e)


#### YT transcript getting below

video_id=input("Enter YT ID here: ")

response = YouTubeTranscriptApi.get_transcript(video_id)


input_transcript = ''

for x in response:
    sentence = x['text']
    input_transcript +=f' {sentence}\n'

# print(output)

# pattern = input('What pattern do you want to use? ')


system_url = "https://raw.githubusercontent.com/tomdowdeswell/patterns/refs/heads/main/extract_learnings/system.md"
# user_url = f"https://raw.githubusercontent.com/tomdowdeswell/patterns/refs/heads/main/{pattern}/user.md"



system_content = fetch_content_from_url(system_url)
# user_file_content = fetch_content_from_url(user_url)


### configure the api request

secret = ""

url = "https://api.openai.com/v1/chat/completions"

headers = {"Authorization": f"Bearer {secret}",
           "Content-Type": "application/json"}

data = {
     "model": "gpt-4o",
     "messages": [
         {"role": "system", "content": system_url},
         {"role": "user", "content": input_transcript}
         ]
   }

payload_data = json.dumps(data)

r = requests.post(url, headers=headers, data=payload_data).json()

content = r['choices'][0]['message']['content']

unsan_title = get_title(video_id)

title = sanitize_filename(unsan_title)

# Save content to a Word document
def save_to_word(content, filename=f"{title}_output.docx"):
    # Create a Word document
    doc = Document()
    doc.add_heading("API Response Content", level=1)
    doc.add_paragraph(content)
    # Save the document
    doc.save(filename)
    print(f"Content successfully saved to {filename}")

# Save the content to a Word document
save_to_word(content)
